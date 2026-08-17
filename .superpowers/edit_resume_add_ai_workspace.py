from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"C:\Users\30829\Desktop\李晨阳-前端偏全栈开发工程师-简历-更新版.docx")
OUTPUT = Path(r"C:\Users\30829\Desktop\ai-collaboration-workspace\李晨阳-前端偏全栈开发工程师-简历-AI协作项目版.docx")


def paragraph_index_by_text(doc: Document, text: str) -> int:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"Paragraph not found: {text}")


def cloned_paragraph(template, texts: list[str]):
    element = deepcopy(template._p)
    text_nodes = element.xpath('.//w:t')
    if len(text_nodes) < len(texts):
        raise ValueError(f"Template has {len(text_nodes)} text nodes, need {len(texts)}")
    for node, value in zip(text_nodes, texts):
        node.text = value
    for node in text_nodes[len(texts):]:
        node.text = ''
    return element


def cloned_link_paragraph(template, document_part, label: str, url: str):
    element = deepcopy(template._p)
    text_runs = [run for run in element.xpath('.//w:r') if run.xpath('./w:t')]
    if not text_runs:
        raise ValueError('Link template has no visible run')
    run_properties = text_runs[0].find(qn('w:rPr'))
    for child in list(element):
        if child.tag != qn('w:pPr'):
            element.remove(child)
    relationship_id = document_part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), relationship_id)
    run = OxmlElement('w:r')
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    text = OxmlElement('w:t')
    text.text = label
    run.append(text)
    hyperlink.append(run)
    element.append(hyperlink)
    return element


doc = Document(SOURCE)
core_heading_index = paragraph_index_by_text(doc, '核心项目')
project_title_index = paragraph_index_by_text(
    doc,
    '政务案件管理系统｜AI 全栈项目  |  React 19 / TypeScript / NestJS / TypeORM / MySQL / Dify / SSE',
)
link_template_index = paragraph_index_by_text(doc, 'github.com/lichenyang5/MiniAppRuntime-Harmony')
bullet_template_index = paragraph_index_by_text(
    doc,
    '前端实现风险预警、事件核查和系统设置三大模块，支持多条件筛选、分页、详情弹窗、风险等级可视化及 Markdown 报告渲染。',
)

anchor = doc.paragraphs[core_heading_index]._p
title_template = doc.paragraphs[project_title_index]
link_template = doc.paragraphs[link_template_index]
bullet_template = doc.paragraphs[bullet_template_index]

elements = [
    cloned_paragraph(
        title_template,
        [
            'AI Collaboration Workspace｜AI 团队协作工作台',
            '  |  React / TypeScript / NestJS / TypeORM / PostgreSQL / Docker',
        ],
    ),
    cloned_link_paragraph(
        link_template,
        doc.part,
        'github.com/lichenyang5/ai-collaboration-workspace',
        'https://github.com/lichenyang5/ai-collaboration-workspace',
    ),
    cloned_paragraph(
        bullet_template,
        ['独立完成团队、项目、成员权限与任务看板的全栈设计，实现登录鉴权、项目管理、任务筛选、编辑、归档恢复及操作活动记录。'],
    ),
    cloned_paragraph(
        bullet_template,
        ['基于 React + TypeScript 构建交互与异步状态管理，使用 NestJS + TypeORM + PostgreSQL 设计 REST API、关系模型和事务化数据访问。'],
    ),
    cloned_paragraph(
        bullet_template,
        ['接入大语言模型，根据项目目标生成结构化任务草案，并实现生成确认、批量持久化、失败反馈及活动记录刷新闭环。'],
    ),
    cloned_paragraph(
        bullet_template,
        ['针对重复邀请、任务更新及团队路由切换设计幂等接口、数据库唯一约束和请求代际隔离；使用 Jest、Vitest 与 Testing Library 建立 137 项自动化测试。'],
    ),
]

for element in elements:
    anchor.addnext(element)
    anchor = element

doc.save(OUTPUT)

reopened = Document(OUTPUT)
full_text = '\n'.join(paragraph.text for paragraph in reopened.paragraphs)
required = [
    'AI Collaboration Workspace｜AI 团队协作工作台',
    'github.com/lichenyang5/ai-collaboration-workspace',
    '137 项自动化测试',
]
for item in required:
    if item not in full_text:
        raise AssertionError(f"Missing output text: {item}")
if full_text.count('AI Collaboration Workspace｜AI 团队协作工作台') != 1:
    raise AssertionError('Project title must appear exactly once')

with ZipFile(OUTPUT) as package:
    document_xml = package.read('word/document.xml')
    rels_xml = package.read('word/_rels/document.xml.rels')
    if b'ai-collaboration-workspace' not in rels_xml:
        raise AssertionError('GitHub hyperlink relationship missing')
    print(f"output={OUTPUT}")
    print(f"paragraphs={len(reopened.paragraphs)}")
    print(f"lastRenderedPageBreak={document_xml.count(b'lastRenderedPageBreak')}")
    print(f"explicitPageBreak={document_xml.count(b'w:type=\"page\"')}")
    print(f"size={OUTPUT.stat().st_size}")
