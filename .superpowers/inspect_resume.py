from pathlib import Path
from docx import Document

path = Path(r"C:\Users\30829\Desktop\李晨阳-前端偏全栈开发工程师-简历-更新版.docx")
doc = Document(path)
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
for i, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.replace("\t", "\\t")
    if text.strip():
        runs = " | ".join(
            f"{run.text!r}[b={run.bold},i={run.italic},size={run.font.size.pt if run.font.size else None},font={run.font.name}]"
            for run in paragraph.runs
        )
        print(f"P{i:03d} style={paragraph.style.name!r} align={paragraph.alignment} text={text!r}")
        print(f"  RUNS {runs}")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else None}")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.replace("\n", " / ") for cell in row.cells]
        print(f"  R{ri}: {cells}")
for si, section in enumerate(doc.sections):
    print(
        f"SECTION {si}: page={section.page_width}x{section.page_height} "
        f"margins={section.top_margin},{section.right_margin},{section.bottom_margin},{section.left_margin}"
    )
