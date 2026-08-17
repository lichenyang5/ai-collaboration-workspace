from pathlib import Path
from zipfile import ZipFile

from docx import Document


source_path = Path(r"C:\Users\30829\Desktop\李晨阳-前端偏全栈开发工程师-简历-更新版.docx")
output_path = Path(r"C:\Users\30829\Desktop\ai-collaboration-workspace\李晨阳-前端偏全栈开发工程师-简历-AI协作项目版.docx")
source = Document(source_path)
output = Document(output_path)

assert len(output.paragraphs) == len(source.paragraphs) + 6
assert len(output.sections) == len(source.sections) == 1
source_section = source.sections[0]
output_section = output.sections[0]
for attribute in (
    'page_width', 'page_height', 'top_margin', 'right_margin',
    'bottom_margin', 'left_margin',
):
    assert getattr(output_section, attribute) == getattr(source_section, attribute), attribute

texts = [paragraph.text for paragraph in output.paragraphs]
heading_index = texts.index('核心项目')
expected_block = [
    'AI Collaboration Workspace｜AI 团队协作工作台  |  React / TypeScript / NestJS / TypeORM / PostgreSQL / Docker',
    'github.com/lichenyang5/ai-collaboration-workspace',
    '独立完成团队、项目、成员权限与任务看板的全栈设计，实现登录鉴权、项目管理、任务筛选、编辑、归档恢复及操作活动记录。',
    '基于 React + TypeScript 构建交互与异步状态管理，使用 NestJS + TypeORM + PostgreSQL 设计 REST API、关系模型和事务化数据访问。',
    '接入大语言模型，根据项目目标生成结构化任务草案，并实现生成确认、批量持久化、失败反馈及活动记录刷新闭环。',
    '针对重复邀请、任务更新及团队路由切换设计幂等接口、数据库唯一约束和请求代际隔离；使用 Jest、Vitest 与 Testing Library 建立 137 项自动化测试。',
]
assert texts[heading_index + 1:heading_index + 7] == expected_block
assert output.paragraphs[heading_index + 1].style.name == 'Normal'
assert output.paragraphs[heading_index + 2].style.name == 'Normal'
for index in range(heading_index + 3, heading_index + 7):
    assert output.paragraphs[index].style.name == 'List Bullet'

with ZipFile(output_path) as package:
    assert package.testzip() is None
    relationships = package.read('word/_rels/document.xml.rels').decode('utf-8')
    assert 'https://github.com/lichenyang5/ai-collaboration-workspace' in relationships

print('STRUCTURAL_QA=PASS')
print(f'PARAGRAPHS={len(output.paragraphs)}')
print('INSERTED_BLOCK:')
for text in expected_block:
    print(text)
